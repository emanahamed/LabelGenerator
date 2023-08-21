import openpyxl
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from docx import Document


def split_address(address):
    max_length = 30
    words = address.split()
    lines = []
    current_line = ""

    for word in words:
        if len(current_line) + len(word) + 1 <= max_length:
            current_line += f"{word} "
        else:
            lines.append(current_line.strip())
            current_line = f"{word} "

    if current_line:
        lines.append(current_line.strip())

    return lines


def create_envelope_label(student_id, student_name, address, pdf_canvas, word_document, x, y, label_width, label_height):
    font_size = 12  # Initial font size
    lines = split_address(address)

    # Dynamically adjust font size to fit 4 lines
    while len(lines) > 4 and font_size > 6:
        font_size -= 1
        pdf_canvas.setFont("Helvetica", font_size)
        lines = split_address(address)

    pdf_canvas.setFont("Helvetica", font_size)
    pdf_canvas.drawString(x + 0.2 * inch, y - 0.3 * inch, f"Parent/Carer of {student_name}")
    pdf_canvas.drawString(x + 0.2 * inch, y - 0.5 * inch, f"Student ID: {student_id}")

    # Ensure exactly 4 lines for the address
    while len(lines) < 4:
        lines.append("")

    for line in lines[:4]:
        pdf_canvas.drawString(x + 0.2 * inch, y - 0.7 * inch, line)
        y -= 12

    paragraph = word_document.add_paragraph()
    paragraph.add_run(f"Parent/Carer of {student_name}\n")
    paragraph.add_run(f"Student ID: {student_id}\n")
    for line in lines[:4]:
        paragraph.add_run(line + "\n")


def create_envelope_labels_pdf_and_docx(file_path):
    wb = openpyxl.load_workbook(file_path)
    sheet = wb.active

    pdf_file = "envelope_labels.pdf"
    pdf_canvas = canvas.Canvas(pdf_file, pagesize=A4)

    docx_file = "envelope_labels.docx"
    docx_document = Document()

    label_width = 2.9 * inch
    label_height = 0.8 * inch
    top_margin = 0.25 * inch
    left_margin = 0.25 * inch
    bottom_margin = 0.25 * inch

    available_width = A4[0] - (2 * left_margin)
    available_height = A4[1] - (top_margin + bottom_margin)

    columns_per_page = int(available_width / label_width)
    rows_per_page = int(available_height / label_height)

    row_count = sheet.max_row
    current_row = 2

    for page in range((row_count - 1) // (columns_per_page * rows_per_page) + 1):
        x, y = left_margin, A4[1] - top_margin

        for row in range(rows_per_page):
            for col in range(columns_per_page):
                if current_row > row_count:
                    break

                student_id = sheet.cell(row=current_row, column=1).value
                student_name = sheet.cell(row=current_row, column=2).value
                address = sheet.cell(row=current_row, column=3).value

                create_envelope_label(
                    student_id, student_name, address, pdf_canvas, docx_document, x, y, label_width, label_height
                )

                x += label_width + 0.2 * inch
                current_row += 1

            x = left_margin
            y -= label_height + 0.2 * inch

        pdf_canvas.showPage()

    pdf_canvas.save()
    print(f"Envelope labels saved as '{pdf_file}'.")
    docx_document.save(docx_file)
    print(f"Envelope labels saved as '{docx_file}'.")


if __name__ == "__main__":
    excel_file_path = "/Users/emanahamed/Desktop/OMZ/AddressLabelPrint/Student Address.xlsx"
    create_envelope_labels_pdf_and_docx(excel_file_path)
